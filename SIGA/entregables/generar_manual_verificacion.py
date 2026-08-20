"""
Genera Manual_Verificacion_SIGCM_en_SIGA.docx

QUE ES
------
El manual de verificacion: como comprobar, dentro del aplicativo SIGA, que lo
que el SIGCM registro llego, y como distinguir el momento del Anexo 3 del
momento del Anexo 4.

POR QUE ES UN SCRIPT Y NO UN .docx SUELTO
-----------------------------------------
Porque las capturas cambian —el aplicativo se actualiza, los datos de prueba se
rehacen— y un Word editado a mano se desactualiza sin que nadie lo note. Aqui el
texto vive en codigo, las capturas en una carpeta, y el documento se regenera.

LAS CAPTURAS
------------
Van en assets_verificacion/ con estos nombres exactos:

    01-menu-modificacion-bienes.png     el menu Programacion > Modificacion de
                                        C.M.N. > Bienes, Servicios y Obras
    02-modificacion-por-area-usuaria.png la grilla con los items del SIGCM
    03-menu-solicitud-modificacion.png  el menu > Solicitud de Modificacion
    04-filtro-area-usuaria.png          el filtro de busqueda del area usuaria
    05-solicitud-aprobada.png           la solicitud con su estado y su sustento

Si alguna falta, el documento se genera igual y en su lugar queda un recuadro
que dice cual falta. Asi el manual sirve desde el primer dia y mejora cuando se
agregan las imagenes, en vez de no existir hasta tenerlas todas.

USO
---
    python generar_manual_verificacion.py
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Falta python-docx.  Instalar con:  python -m pip install python-docx pillow")

try:
    from PIL import Image
    TIENE_PIL = True
except ImportError:
    TIENE_PIL = False

AQUI = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(AQUI, "assets_verificacion")
SALIDA = os.path.join(AQUI, "Manual_Verificacion_SIGCM_en_SIGA.docx")

ANCHO_UTIL_CM = 16.0          # A4 vertical con margenes de 2,5 cm
AZUL = RGBColor(0x1F, 0x3A, 0x93)
GRIS = RGBColor(0x55, 0x5555 // 0x100, 0x55)
ROJO = RGBColor(0xB0, 0x1C, 0x1C)


# --------------------------------------------------------------------------- #
# Utilidades de formato                                                       #
# --------------------------------------------------------------------------- #

def sombrear(celda, color_hex):
    """Fondo de celda. python-docx no lo expone; hay que bajar al XML."""
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc.append(shd)


def parrafo(doc, texto="", negrita=False, tam=10.5, color=None,
            espacio_antes=0, espacio_despues=6, alineacion=None):
    p = doc.add_paragraph()
    if alineacion is not None:
        p.alignment = alineacion
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    if texto:
        r = p.add_run(texto)
        r.bold = negrita
        r.font.size = Pt(tam)
        if color is not None:
            r.font.color.rgb = color
    return p


def mixto(doc, trozos, tam=10.5, espacio_despues=6):
    """Un parrafo con partes en negrita y partes normales: [(texto, negrita)]."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(espacio_despues)
    for texto, negrita in trozos:
        r = p.add_run(texto)
        r.bold = negrita
        r.font.size = Pt(tam)
    return p


def codigo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(texto)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def aviso(doc, titulo, cuerpo, color_fondo="FFF4E5", color_texto=ROJO):
    """Recuadro de una celda, para lo que no se puede pasar por alto."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, color_fondo)
    c.width = Cm(ANCHO_UTIL_CM)

    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = color_texto

    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(cuerpo)
    r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def captura(doc, nombre_archivo, pie):
    """Inserta la captura, o un recuadro que dice cual falta."""
    ruta = os.path.join(ASSETS, nombre_archivo)

    if os.path.exists(ruta):
        ancho = Cm(ANCHO_UTIL_CM)
        if TIENE_PIL:
            with Image.open(ruta) as im:
                w, h = im.size
            # Una captura muy apaisada no necesita todo el ancho si al reducirla
            # el texto queda ilegible; se deja igual, pero se avisa en consola.
            if w / float(h) > 3.2:
                print("   [aviso] %s es muy apaisada (%dx%d): revisar legibilidad"
                      % (nombre_archivo, w, h))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(ruta, width=ancho)
        estado = "ok"
    else:
        t = doc.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        sombrear(c, "F0F0F0")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run("[ Falta la captura:  assets_verificacion/%s ]" % nombre_archivo)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = ROJO
        estado = "falta"

    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_after = Pt(12)
    rr = pp.add_run(pie)
    rr.italic = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return estado


def tabla(doc, encabezados, filas, anchos_cm=None, resaltar_col=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, texto in enumerate(encabezados):
        c = t.rows[0].cells[i]
        sombrear(c, "1F3A93")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for n, fila in enumerate(filas):
        celdas = t.add_row().cells
        for i, texto in enumerate(fila):
            c = celdas[i]
            if n % 2 == 1:
                sombrear(c, "F4F6FB")
            if resaltar_col is not None and i == resaltar_col:
                sombrear(c, "E8F0FE")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            negrita = texto.startswith("*") and texto.endswith("*")
            r = p.add_run(texto.strip("*"))
            r.bold = negrita
            r.font.size = Pt(9.5)

    if anchos_cm:
        for fila in t.rows:
            for i, ancho in enumerate(anchos_cm):
                fila.cells[i].width = Cm(ancho)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# --------------------------------------------------------------------------- #
# El documento                                                                #
# --------------------------------------------------------------------------- #

def construir():
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    faltantes = []

    # ---- Portada --------------------------------------------------------- #
    parrafo(doc, "AUTORIDAD NACIONAL DE INFRAESTRUCTURA — ANIN",
            negrita=True, tam=11, color=AZUL, espacio_antes=36,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=2)
    parrafo(doc, "Sistema de Gestión de Contrataciones Menores (SIGCM)",
            tam=10.5, alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=28)

    parrafo(doc, "Manual de verificación del registro en SIGA",
            negrita=True, tam=22, color=AZUL,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=6)
    parrafo(doc, "Cómo comprobar en el aplicativo SIGA lo que el SIGCM registró, "
                 "y cómo distinguir el Anexo 3 del Anexo 4",
            tam=12, alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=36)

    tabla(doc,
          ["", ""],
          [["*Módulo*", "Logística — Programación del C.M.N."],
           ["*Aplica a*", "Cuadro Multianual de Necesidades, ejercicio 2026"],
           ["*Base de referencia*", "SIGA_1750 (copia local de trabajo)"],
           ["*Directiva*", "N.° 0007-2025-EF/54.01"],
           ["*Versión*", "1.0 — 20 de agosto de 2026"]],
          anchos_cm=[4.5, 11.5])

    doc.add_page_break()

    # ---- 1. Para qué sirve ------------------------------------------------ #
    parrafo(doc, "1.  Para qué sirve este manual", negrita=True, tam=15,
            color=AZUL, espacio_despues=8)

    parrafo(doc,
            "El SIGCM lleva el trámite de la modificación del Cuadro Multianual de "
            "Necesidades —registro, firmas, derivaciones— y en dos momentos concretos "
            "escribe en la base de SIGA. Este manual explica cómo confirmar, desde el "
            "propio aplicativo SIGA, que esa escritura ocurrió y en qué momento del "
            "trámite está cada ítem.")

    parrafo(doc,
            "Está escrito para responder una sola pregunta, que se ha hecho varias "
            "veces: «registré desde el SIGCM, ¿dónde lo veo en SIGA?».")

    aviso(doc,
          "La respuesta corta",
          "Logística → Programación → Modificación de C.M.N. → Bienes, Servicios y Obras.\n"
          "No en «Programación del C.M.N.», y no en «Demanda Adicional — Identificación».",
          color_fondo="E8F0FE", color_texto=AZUL)

    parrafo(doc, "1.1  Por qué no es donde uno cree", negrita=True, tam=12.5,
            color=AZUL, espacio_antes=10, espacio_despues=6)

    parrafo(doc,
            "El módulo Logística separa las tres etapas del CMN en ramas distintas del "
            "menú, y cada rama lee una tabla distinta. Ahí está todo el enredo:")

    tabla(doc,
          ["Rama del menú", "Qué muestra", "Tabla"],
          [["Programación del C.M.N.", "El cuadro tal como se formuló y aprobó. Es una foto que ya no cambia.", "SIG_CUADRO_NECESIDAD_DET"],
           ["*Modificación de C.M.N.*", "*Lo que se modifica durante el año. Aquí está lo del SIGCM.*", "*SIG_CUADRO_MODIFICADO_DET*"],
           ["Consolidado del C.M.N.", "El consolidado y el PAAC. Es un proceso por lotes de SIGA.", "SIG_CUADRO_MODIFICADO_CMN"]],
          anchos_cm=[4.2, 7.3, 4.5])

    aviso(doc,
          "«Demanda Adicional» no sirve para esto",
          "Es un botón dentro de la pantalla de formulación, no la vista del cuadro "
          "modificado. Además exige la habilitación flag_da_aprob, que sólo otorga "
          "Abastecimiento desde SIGA, y muestra únicamente la demanda adicional EN "
          "CURSO: en cuanto el Anexo 4 se firma, el ítem desaparece de ahí. Es decir, "
          "en el único momento en que uno quiere confirmar que todo salió bien, esa "
          "pantalla está vacía.")

    doc.add_page_break()

    # ---- 2. Los dos momentos --------------------------------------------- #
    parrafo(doc, "2.  Los dos momentos de escritura", negrita=True, tam=15,
            color=AZUL, espacio_despues=8)

    parrafo(doc,
            "El SIGCM no escribe en SIGA una vez, sino dos. Confundirlas lleva a dar "
            "por disponible un ítem que todavía no lo está.")

    mixto(doc, [("Primero, la firma del jefe de Abastecimiento sobre el ", False),
                ("Anexo 3", True),
                (" registra el ítem: existe en SIGA, pero SIGA no lo deja pedir. "
                 "Después, la firma sobre el ", False),
                ("Anexo 4", True),
                (" lo habilita.", False)])

    tabla(doc,
          ["", "Después del Anexo 3", "Después del Anexo 4"],
          [["*Acción en el SIGCM*", "El jefe de Abastecimiento firma el Anexo 3", "El jefe de Abastecimiento firma el Anexo 4"],
           ["*Estado en pantalla*", "*V.B. Jefe*", "*Aprobado*"],
           ["Fecha de V.B. Jefe", "con fecha", "con fecha"],
           ["*Fecha de Aprobación*", "*vacía*", "*con fecha*"],
           ["Incl / Excl del ítem", "I", "I  (no cambia)"],
           ["ESTADO del detalle", "I", "I  (no cambia)"],
           ["FLAG_MODIFICADO", "1", "0"],
           ["*MOTIVO_SOLICITUD*", "*1*", "*0*"],
           ["*¿Se puede pedir el ítem?*", "*NO*", "*SÍ*"]],
          anchos_cm=[5.0, 5.5, 5.5])

    mixto(doc, [("La columna que decide es ", False), ("MOTIVO_SOLICITUD", True),
                (". El selector de ítems de un requerimiento exige "
                 "MOTIVO_SOLICITUD IN ('0','3'). Mientras valga 1, el ítem existe "
                 "pero no es pedible. Eso no es un defecto: es exactamente lo que "
                 "significa «Anexo 3 firmado, Anexo 4 pendiente».", False)])

    codigo(doc,
           "Anexo 3 firmado  ->  el item EXISTE en SIGA, pero no es pedible   (V.B. Jefe)\n"
           "Anexo 4 firmado  ->  el item queda HABILITADO                     (Aprobado)")

    aviso(doc,
          "El Anexo 4 no consolida",
          "Consolidar es generar el árbol del PAAC, un proceso por lotes de SIGA. "
          "De las 5 837 inclusiones aprobadas de 2026, sólo 4 406 están consolidadas. "
          "Buscar lo nuestro en «Consolidado del C.M.N.» y no encontrarlo es lo "
          "esperado, no un error.")

    doc.add_page_break()

    # ---- 3. Verificar el Anexo 3 ------------------------------------------ #
    parrafo(doc, "3.  Verificación después del Anexo 3", negrita=True, tam=15,
            color=AZUL, espacio_despues=8)

    parrafo(doc,
            "Se hace en cuanto el jefe de Abastecimiento firma el Anexo 3 en el SIGCM "
            "y el worker de integración drena la cola.")

    parrafo(doc, "Paso 1 — Abrir la pantalla del cuadro modificado",
            negrita=True, tam=12.5, color=AZUL, espacio_antes=8, espacio_despues=6)

    codigo(doc, "Logistica  >  Programacion  >  Modificacion de C.M.N.  >  Bienes, Servicios y Obras")

    faltantes.append(captura(doc, "01-menu-modificacion-bienes.png",
                             "Figura 1. La ruta del menú. Es «Modificación de C.M.N.», "
                             "no «Programación del C.M.N.»."))

    parrafo(doc, "Paso 2 — Filtrar", negrita=True, tam=12.5, color=AZUL,
            espacio_antes=6, espacio_despues=6)

    tabla(doc,
          ["Campo", "Qué poner", "Si se omite"],
          [["Año", "2026", "No trae nada"],
           ["Área Usuaria", "El centro de costo, p. ej. 01.07.05.03. El binocular abre el buscador.", "No trae nada"],
           ["*Tipo*", "*Servicio (o Bien, según lo registrado)*", "*La grilla queda vacía*"]],
          anchos_cm=[3.0, 9.0, 4.0])

    parrafo(doc, "Paso 3 — Encontrar la fila", negrita=True, tam=12.5,
            color=AZUL, espacio_antes=6, espacio_despues=6)

    parrafo(doc,
            "La grilla se agrupa por Actividad Operativa. Los ítems que registró el "
            "SIGCM aparecen junto a los que registraron usuarios reales; se "
            "distinguen por el código de ítem, por el monto y por la columna "
            "Incl / Excl.")

    faltantes.append(captura(doc, "02-modificacion-por-area-usuaria.png",
                             "Figura 2. Modificación de C.M.N. por Área Usuaria. "
                             "Recuadradas, las dos líneas que registró el SIGCM: "
                             "«I» en Incl/Excl y 4 000,00 en el monto de 2026."))

    tabla(doc,
          ["Columna", "Qué debe decir"],
          [["*Incl / Excl*", "*I para una inclusión, E para una exclusión*"],
           ["Item", "El código de catálogo, p. ej. 210100010579"],
           ["Descripción", "La del ítem de catálogo"],
           ["Clasificador", "El que se envió, p. ej. 2.3. 2  9. 1  1"],
           ["Meta", "p. ej. 0015"],
           ["2026 → Mnto. Total", "El monto de la inclusión"]],
          anchos_cm=[4.5, 11.5])

    aviso(doc,
          "Cómo encontrarla rápido",
          "Los casos de prueba se cargan de setiembre a diciembre a propósito. Lo que "
          "las áreas ya tenían está en meses anteriores, así que la línea nueva salta "
          "a la vista.",
          color_fondo="E8F5E9", color_texto=RGBColor(0x1B, 0x5E, 0x20))

    doc.add_page_break()

    parrafo(doc, "Paso 4 — Comprobar el estado de la solicitud",
            negrita=True, tam=12.5, color=AZUL, espacio_despues=6)

    codigo(doc, "Logistica  >  Programacion  >  Modificacion de C.M.N.  >  Solicitud de Modificacion")

    faltantes.append(captura(doc, "03-menu-solicitud-modificacion.png",
                             "Figura 3. La segunda pantalla que se usa: Solicitud de "
                             "Modificación."))

    parrafo(doc,
            "Filtrar por Año 2026 y marcar «Todos» en Mes y en Nº Solicitud. El "
            "binocular del Área Usuaria abre el filtro de búsqueda; se puede buscar "
            "por Descripción.")

    faltantes.append(captura(doc, "04-filtro-area-usuaria.png",
                             "Figura 4. El filtro de búsqueda del área usuaria. "
                             "Ojo: 01.07.05 y 01.07.05.03 se llaman casi igual."))

    aviso(doc,
          "Dos áreas con el mismo nombre",
          "01.07.05 «OFICINA DE TECNOLOGÍAS DE LA INFORMACIÓN» y 01.07.05.03 "
          "«OFICINA DE TECNOLOGÍAS DE LA INFORMACIÓN.» se distinguen sólo por el "
          "punto final y por el código. La que usamos es 01.07.05.03.")

    mixto(doc, [("En la lista de la izquierda, la solicitud del Anexo 3 recién "
                 "firmado aparece con estado ", False),
                ("V.B. Jefe", True),
                (". Al abrirla, el ", False), ("Sustento", True),
                (" dice ", False), ("INCLUSION CMN INTEGRADA DESDE SIGCM", True),
                (" y la ", False), ("Fecha de Aprobación está vacía", True),
                (".", False)])

    doc.add_page_break()

    # ---- 4. Verificar el Anexo 4 ------------------------------------------ #
    parrafo(doc, "4.  Verificación después del Anexo 4", negrita=True, tam=15,
            color=AZUL, espacio_despues=8)

    parrafo(doc,
            "Se hace cuando el jefe de Abastecimiento firma el Anexo 4. Si el Anexo 4 "
            "agrupó Anexos 3 de varias áreas usuarias, hay que repetir la "
            "comprobación con cada una: SIGA recibe una aprobación por área, no una "
            "sola por el Anexo 4.")

    parrafo(doc, "Qué cambia respecto del paso anterior", negrita=True,
            tam=12.5, color=AZUL, espacio_antes=6, espacio_despues=6)

    tabla(doc,
          ["Pantalla", "Antes (Anexo 3)", "Ahora (Anexo 4)"],
          [["Bienes, Servicios y Obras", "El ítem con I en Incl/Excl", "*Igual: el ítem no cambia de aspecto*"],
           ["*Solicitud de Modificación*", "*Estado: V.B. Jefe*", "*Estado: Aprobado*"],
           ["Solicitud de Modificación", "Fecha de Aprobación vacía", "*Fecha de Aprobación con fecha*"]],
          anchos_cm=[5.0, 5.5, 5.5])

    aviso(doc,
          "En la grilla del cuadro el ítem NO cambia",
          "Es lo que más confunde. La aprobación del Anexo 4 no altera la fila del "
          "cuadro modificado: cambia MOTIVO_SOLICITUD de 1 a 0, que la pantalla no "
          "muestra. La confirmación visual está en la pantalla de Solicitud de "
          "Modificación, en el estado y en la Fecha de Aprobación.")

    faltantes.append(captura(doc, "05-solicitud-aprobada.png",
                             "Figura 5. La solicitud ya aprobada: estado «Aprobado», "
                             "Fecha de Aprobación con fecha y el sustento «INCLUSION "
                             "CMN INTEGRADA DESDE SIGCM»."))

    parrafo(doc, "La comprobación definitiva", negrita=True, tam=12.5,
            color=AZUL, espacio_antes=6, espacio_despues=6)

    parrafo(doc,
            "El ítem está realmente habilitado cuando aparece en el selector de ítems "
            "de un requerimiento. Ése es el efecto que se buscaba y el que el área "
            "usuaria va a notar.")

    doc.add_page_break()

    # ---- 5. Verificacion por consulta ------------------------------------- #
    parrafo(doc, "5.  Verificación por consulta", negrita=True, tam=15,
            color=AZUL, espacio_despues=8)

    parrafo(doc,
            "Cuando la pantalla no basta —o para comprobar varias áreas de una vez— "
            "se consulta directamente. Sobre la base SIGA_1750:")

    codigo(doc,
           "-- En que momento esta cada solicitud del SIGCM\n"
           "SELECT SEC_SOL_MOD, CENTRO_COSTO, ESTADO, FECHA, LEFT(GLOSA,45) AS GLOSA\n"
           "  FROM dbo.SIG_SOLICITUD_MODIFICACION\n"
           " WHERE ANNO_EJEC = 2026 AND SEC_EJEC = 1750\n"
           "   AND GLOSA LIKE '%SIGCM%' COLLATE Modern_Spanish_CI_AI\n"
           " ORDER BY SEC_SOL_MOD DESC;\n"
           "\n"
           "-- ESTADO 2 = V.B. Jefe (Anexo 3)     ESTADO 3 = Aprobado (Anexo 4)")

    codigo(doc,
           "-- Como quedo el item\n"
           "SELECT SEC_CUADRO, SEC_ITEM, ANNO_PROG, ESTADO, PROCEDENCIA,\n"
           "       FLAG_MODIFICADO, MOTIVO_SOLICITUD, CANT_TOTAL, MNTO_TOTAL\n"
           "  FROM dbo.SIG_CUADRO_MODIFICADO_DET\n"
           " WHERE ANNO_EJEC = 2026 AND SEC_EJEC = 1750\n"
           "   AND CENTRO_COSTO = '01.07.05.03'\n"
           " ORDER BY SEC_ITEM DESC, ANNO_PROG;\n"
           "\n"
           "-- MOTIVO_SOLICITUD 1 = Anexo 3       0 = Anexo 4")

    parrafo(doc, "Y sobre DBSIGCM, el puente entre ambos mundos:",
            espacio_antes=4)

    codigo(doc,
           "SELECT s.Codigo, m.SecCuadro, m.SecItem, m.SecSolicitud\n"
           "  FROM integracion.MapeoCmn AS m\n"
           "  JOIN cmn.Solicitud AS s ON s.IdSolicitud = m.IdSolicitud\n"
           " ORDER BY s.Codigo DESC;")

    # ---- 6. Areas de prueba ----------------------------------------------- #
    parrafo(doc, "6.  Áreas usuarias aptas para verificar", negrita=True,
            tam=15, color=AZUL, espacio_antes=14, espacio_despues=8)

    parrafo(doc,
            "No cualquier área sirve. Hacen falta tres condiciones, y ninguna la "
            "valida la base: si falta alguna, la escritura funciona igual pero el "
            "registro queda invisible o el techo lo rechaza.")

    tabla(doc,
          ["Condición", "Dónde se comprueba"],
          [["El cuadro del centro está en estado 4", "SIG_CUADRO_X_CENTRO.estado = '4'"],
           ["Hay techo libre en esa meta y clasificador", "SIG_TECHO_PRESUPUESTO frente a SIG_CUADRO_MODIFICADO_DET"],
           ["El centro tiene tarea activa", "SIG_CENTRO_COSTO_TAREA.estado = 'A'"]],
          anchos_cm=[7.0, 9.0])

    parrafo(doc, "Áreas verificadas al 20 de agosto de 2026:",
            espacio_antes=6, espacio_despues=6)

    tabla(doc,
          ["Centro", "Área", "Meta", "Clasificador", "Libre 2026", "Ítems"],
          [["01.07.05.03", "OTI", "15", "2.3. 2  9. 1  1", "123 994", "18"],
           ["01.07.05.01", "UDS", "11", "2.3. 2  5. 1 99", "360 731", "36"],
           ["01.07.05.02", "US", "14", "2.3. 2  9. 1  1", "102 999", "192"],
           ["01.07.04", "ORH", "18", "2.3. 2  7. 3  1", "68 700", "147"]],
          anchos_cm=[3.0, 1.8, 1.5, 4.2, 2.8, 1.7])

    aviso(doc,
          "01.01 (JEFATURA) no sirve para verificar en el aplicativo",
          "Su cuadro está en estado 6, no 4. Se puede usar para ejercitar el SIGCM, "
          "pero lo que se registre ahí no se verá nunca en SIGA.")

    parrafo(doc,
            "El estado 6 se muestra en SIGA como «C.C.M.N.» y el 4 como "
            "«Consolidación y Aprobación». Suena al revés de lo que uno espera, y ya "
            "se leyó al revés una vez.", espacio_antes=4)

    # ---- 7. Referencias --------------------------------------------------- #
    parrafo(doc, "7.  Documentos relacionados", negrita=True, tam=15,
            color=AZUL, espacio_antes=14, espacio_despues=8)

    tabla(doc,
          ["Documento", "Qué contiene"],
          [["Proyecto/SIGA_APLICATIVO.md", "Este mismo recorrido en texto plano, para consulta rápida"],
           ["Proyecto/CONTEXTO.md", "El contexto del proyecto y la bitácora de iteraciones"],
           ["SIGA/integracion/ANALISIS_CMN.md", "El análisis completo de la integración, con la evidencia"],
           ["SIGA/integracion/captura_siga_xe.sql", "Extended Events, para ver el SQL real que envía el aplicativo"],
           ["db/90_pruebas/S904__casos_anexo4_multiple.sql", "Deja cuatro Anexos 3 listos para recorrer el flujo"]],
          anchos_cm=[7.0, 9.0])

    doc.save(SALIDA)
    return faltantes


if __name__ == "__main__":
    if not os.path.isdir(ASSETS):
        os.makedirs(ASSETS)
        print("Creada la carpeta de capturas: %s" % ASSETS)

    estados = construir()
    faltan = estados.count("falta")

    print("Generado: %s" % SALIDA)
    if faltan:
        print("  %d captura(s) sin archivo. El documento las marca en su lugar." % faltan)
        print("  Guardarlas en %s y volver a ejecutar este script." % ASSETS)
    else:
        print("  Las 5 capturas quedaron incrustadas.")
